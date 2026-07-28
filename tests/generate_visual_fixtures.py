from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from openpyxl import Workbook


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    output = Path(args.output); output.mkdir(parents=True, exist_ok=True)
    document = Document(); document.add_heading("视觉回归夹具", 0); document.add_paragraph("用于验证 Word 页面渲染链路。")
    document.save(output / "fixture.docx")
    workbook = Workbook(); sheet = workbook.active; sheet.title = "三线表"; sheet.append(["变量", "估计值", "95% CI"]); sheet.append(["示例", "1.20", "0.80–1.60"])
    workbook.save(output / "fixture.xlsx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
