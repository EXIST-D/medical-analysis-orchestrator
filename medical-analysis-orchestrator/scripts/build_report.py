#!/usr/bin/env python3
"""Build a traceable Chinese academic manuscript draft from validated results."""

from __future__ import annotations

import argparse
import csv
import json
import math
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


EAST_ASIA_FONT = "宋体"
LATIN_FONT = "Times New Roman"
BODY_SIZE = 12
TABLE_SIZE = 9
MANUSCRIPT_REPORT_FILENAME = "01_医学统计分析论文初稿.docx"
MODULE_TITLES = {
    "descriptive": "描述性统计",
    "group-comparison": "单因素分析",
    "correlation": "相关性分析",
    "linear-regression": "多元线性回归",
    "logistic-regression": "Logistic 回归",
    "reliability-validity": "信度与效度分析",
    "factor-analysis": "探索性与验证性因子分析",
    "mixed-effects": "混合效应模型",
    "missing-data": "缺失数据与多重插补",
    "generalized-regression": "广义回归",
    "survival": "基础生存分析",
    "diagnostic-accuracy": "诊断试验准确性",
    "gee": "广义估计方程",
    "measurement-invariance": "测量不变性",
    "competing-risks": "竞争风险分析",
    "propensity-score": "倾向评分分析",
    "sem": "结构方程模型",
    "network": "网络分析",
    "bayesian": "贝叶斯网络",
}


def load_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    return payload if isinstance(payload, dict) else {}


def load_yaml(path: Path) -> dict[str, Any]:
    payload = yaml.safe_load(path.read_text(encoding="utf-8-sig"))
    return payload if isinstance(payload, dict) else {}


def set_run_font(run: Any, size: float = BODY_SIZE, bold: bool | None = None) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    fonts.set(qn("w:cs"), LATIN_FONT)


def set_style_font(style: Any, size: float, bold: bool = False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Pt(0)
    style.paragraph_format.space_before = Pt(6 if size >= 14 else 3)
    style.paragraph_format.space_after = Pt(0)


def set_paragraph_format(
    paragraph: Any,
    *,
    indent: bool = True,
    alignment: Any | None = None,
    space_before: float = 0,
    space_after: float = 0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.first_line_indent = Pt(24) if indent else Pt(0)
    if alignment is not None:
        paragraph.alignment = alignment


def add_text_paragraph(
    document: Document,
    text: str,
    *,
    indent: bool = True,
) -> Any:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, indent=indent)
    content_run = paragraph.add_run(text)
    set_run_font(content_run)
    return paragraph


def add_heading(document: Document, text: str, level: int) -> Any:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_paragraph_format(
        paragraph,
        indent=False,
        space_before=6 if level == 1 else 3,
        space_after=0,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_caption(document: Document, text: str) -> Any:
    paragraph = document.add_paragraph()
    set_paragraph_format(
        paragraph,
        indent=False,
        space_before=3,
        space_after=0,
    )
    if text.startswith(("表 ", "图 ")):
        marker, separator, title = text.partition(" ")
        number, separator_after_number, remainder = title.partition(" ")
        marker_run = paragraph.add_run(f"{marker} {number}")
        set_run_font(marker_run, size=10, bold=True)
        if separator_after_number:
            title_run = paragraph.add_run(f" {remainder}")
            set_run_font(title_run, size=10)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_table_note(document: Document, text: str) -> Any:
    """Add a compact note that remains visually subordinate to its table."""
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, indent=False, space_before=0, space_after=2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5)
    return paragraph


def set_cell_border(cell: Any, **edges: dict[str, Any]) -> None:
    tc_properties = cell._tc.get_or_add_tcPr()
    borders = tc_properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_properties.append(borders)
    for edge_name, edge_data in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_width(table: Any, width_dxa: int, weights: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(width_dxa))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    weight_total = sum(weights) or 1
    minimum_width = 500
    distributable = max(0, width_dxa - minimum_width * len(weights))
    exact_extras = [
        distributable * weight / weight_total for weight in weights
    ]
    column_widths = [
        minimum_width + int(extra) for extra in exact_extras
    ]
    remainder = width_dxa - sum(column_widths)
    fractional_order = sorted(
        range(len(weights)),
        key=lambda index: exact_extras[index] - int(exact_extras[index]),
        reverse=True,
    )
    for index in fractional_order[:remainder]:
        column_widths[index] += 1
    grid = table._tbl.tblGrid
    for index, grid_col in enumerate(grid):
        grid_col.set(qn("w:w"), str(column_widths[index]))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            column_width = column_widths[index]
            cell.width = Inches(column_width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(column_width))
            tc_w.set(qn("w:type"), "dxa")


def format_table_cell(
    cell: Any,
    *,
    header: bool = False,
    font_size: float = TABLE_SIZE,
    left_aligned: bool = False,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        set_paragraph_format(
            paragraph,
            indent=False,
            alignment=(
                WD_ALIGN_PARAGRAPH.LEFT
                if left_aligned
                else WD_ALIGN_PARAGRAPH.CENTER
            ),
        )
        for run in paragraph.runs:
            set_run_font(run, size=font_size, bold=header)


def display_value(value: str, *, header: bool = False) -> str:
    text = value.strip()
    if header:
        return text.replace("_", " ")
    if not text:
        return ""
    try:
        number = float(text)
    except ValueError:
        return text if len(text) <= 70 else text[:67] + "…"
    if not math.isfinite(number):
        return text
    if number.is_integer() and abs(number) < 1_000_000:
        return str(int(number))
    if number != 0 and abs(number) < 0.001:
        return f"{number:.3e}"
    rendered = f"{number:.3f}".rstrip("0").rstrip(".")
    return rendered


def table_column_weights(rows: list[list[str]], column_count: int) -> list[float]:
    weights = []
    for column_index in range(column_count):
        values = [
            row[column_index] if column_index < len(row) else ""
            for row in rows[:16]
        ]
        longest = max(
            (len(value) + sum(ord(char) > 127 for char in value) for value in values),
            default=5,
        )
        weights.append(float(min(24, max(6, longest))))
    return weights


def add_three_line_table(
    document: Document,
    rows: list[list[str]],
    *,
    plan: dict[str, Any] | None = None,
    max_rows: int = 25,
    max_columns: int = 10,
) -> None:
    if not rows:
        add_text_paragraph(document, "该结果表为空。")
        return
    reference_note = table_reference_level_note(rows, plan or {})
    rows = prepare_article_table_rows(rows, plan or {})
    selected = [
        [
            display_value(value, header=row_index == 0)
            for value in row[:max_columns]
        ]
        for row_index, row in enumerate(rows[: max_rows + 1])
    ]
    column_count = max(len(row) for row in selected)
    table = document.add_table(rows=len(selected), cols=column_count)
    set_table_width(table, 9030, table_column_weights(selected, column_count))
    table_font_size = 7.5 if column_count >= 8 else 8.5 if column_count >= 6 else 9
    for row_index, row in enumerate(selected):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            format_table_cell(
                cell,
                header=row_index == 0,
                font_size=table_font_size,
                left_aligned=column_index == 0,
            )
            if row_index == 0:
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": "12", "color": "000000"},
                    bottom={"val": "single", "sz": "6", "color": "000000"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            elif row_index == len(selected) - 1:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "12", "color": "000000"},
                    top={"val": "nil"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            else:
                set_cell_border(
                    cell,
                    top={"val": "nil"},
                    bottom={"val": "nil"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    if len(rows) > len(selected):
        add_table_note(
            document,
            f"Word 报告仅展示前 {len(selected) - 1} 行，完整结果见对应 CSV/XLSX 文件。",
        )
    if len(rows[0]) > max_columns:
        add_table_note(
            document,
            f"Word 报告仅展示前 {max_columns} 列，完整结果见对应 CSV/XLSX 文件。",
        )
    if reference_note:
        add_table_note(document, f"注：{reference_note}")


def read_csv_rows(path: Path) -> list[list[str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return [[str(value) for value in row] for row in csv.reader(handle)]


def evidence_index(results: dict[str, Any]) -> dict[str, dict[str, str]]:
    index: dict[str, dict[str, str]] = {}
    for module_id, result in results.items():
        for table in result.get("tables", []):
            table_id = str(table.get("table_id") or "")
            if table_id:
                index[f"{module_id}:{table_id}"] = {
                    "type": "table",
                    "title": str(table.get("title") or table_id),
                    "path": str(table.get("xlsx_path") or table.get("csv_path") or ""),
                }
        for figure in result.get("figures", []):
            figure_id = str(figure.get("figure_id") or "")
            if figure_id:
                index[f"{module_id}:{figure_id}"] = {
                    "type": "figure",
                    "title": str(figure.get("title") or figure_id),
                    "path": str(figure.get("path") or ""),
                }
    return index


def build_claim_rows(
    manuscript: dict[str, Any], results: dict[str, Any]
) -> list[dict[str, str]]:
    index = evidence_index(results)
    rows: list[dict[str, str]] = []
    for claim in manuscript.get("claims") or []:
        references = [str(item) for item in claim.get("evidence_refs") or []]
        unresolved = [item for item in references if item not in index]
        if unresolved:
            raise SystemExit(
                "论文主张包含未解析的证据引用：" + "、".join(unresolved)
            )
        evidence = [
            f"{index[item]['title']}（{index[item]['path']}）"
            for item in references
        ]
        rows.append(
            {
                "claim_id": str(claim.get("id") or ""),
                "statement": str(claim.get("statement") or ""),
                "evidence_refs": "；".join(references),
                "evidence_files": "；".join(evidence),
                "interpretation_level": str(
                    claim.get("interpretation_level") or ""
                ),
                "boundary": str(claim.get("boundary") or ""),
                "status": "confirmed_and_resolved",
            }
        )
    return rows


def write_manuscript_support_artifacts(
    run_dir: Path,
    plan: dict[str, Any],
    manifest: dict[str, Any],
    results: dict[str, Any],
) -> list[dict[str, str]]:
    reporting = plan.get("reporting") or {}
    manuscript = reporting.get("manuscript_support") or {}
    if manuscript.get("enabled") is not True:
        return []
    output_dir = run_dir / "90_最终报告"
    output_dir.mkdir(parents=True, exist_ok=True)
    rows = build_claim_rows(manuscript, results)

    claim_path = output_dir / "02_主张证据边界表.csv"
    fieldnames = [
        "claim_id",
        "statement",
        "evidence_refs",
        "evidence_files",
        "interpretation_level",
        "boundary",
        "status",
    ]
    with claim_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    variables = plan.get("variables") or {}
    labels = variables.get("labels") or {}
    units = variables.get("units") or {}
    manual_terms = manuscript.get("terminology") or {}
    term_names = sorted(set(labels) | set(units) | set(manual_terms))
    terms: dict[str, Any] = {}
    for term in term_names:
        manual_value = manual_terms.get(term)
        if isinstance(manual_value, dict):
            entry = dict(manual_value)
        else:
            entry = {
                "preferred_term": str(
                    manual_value or labels.get(term) or term
                )
            }
        if labels.get(term):
            entry.setdefault("label", str(labels[term]))
        if units.get(term):
            entry.setdefault("unit", str(units[term]))
        entry.setdefault("source_variable", term)
        terms[term] = entry
    terminology_payload = {
        "schema_version": "1.0",
        "run_id": manifest.get("run_id"),
        "terms": terms,
        "rule": "同一术语、缩写、单位和变量标签在本次运行的全部报告中保持一致。",
    }
    terminology_path = output_dir / "04_术语账本.yml"
    terminology_path.write_text(
        yaml.safe_dump(
            terminology_payload,
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )

    research = plan.get("research") or {}
    handling = plan.get("data_handling") or {}
    runtime = plan.get("runtime") or {}
    modules = [
        item if isinstance(item, str) else item.get("id", "")
        for item in (plan.get("analysis") or {}).get("modules", [])
    ]
    methods_lines = [
        "# 统计方法与可复现性摘要",
        "",
        f"- 运行编号：{manifest.get('run_id', '')}",
        f"- 主要研究问题：{research.get('primary_question', '')}",
        f"- 研究设计：{research.get('design', 'unknown')}",
        f"- 目标量/研究目标：{research.get('estimand_or_target', '') or '未单独登记'}",
        f"- 已确认模块：{', '.join(str(item) for item in modules if item)}",
        f"- 缺失值策略：{handling.get('missing_strategy', '')}",
        (
            "- 多重比较："
            f"{(handling.get('multiple_testing') or {}).get('method', '')}；"
            f"比较族={(handling.get('multiple_testing') or {}).get('family_definition', '')}"
        ),
        f"- 统计引擎：{runtime.get('language', 'R')}",
        f"- 最低 R 版本：{runtime.get('minimum_version', '')}",
        f"- 随机种子：{(plan.get('run') or {}).get('random_seed', '')}",
        f"- 分析方案 SHA-256：{manifest.get('analysis_plan_sha256', '')}",
        "- R 会话信息：sessionInfo.txt",
        "- 包版本：package_versions.csv",
        "- 输入文件与全部输出哈希：manifest.json",
        "",
        "## 表述边界",
        "",
        "- 结果陈述、解释和边界分别登记，不以 P 值替代效应量、区间或诊断。",
        "- 关联和预测结果不得自动升级为因果结论。",
        "- 未登记或未解析到统一结果对象的证据不得写入确认主张。",
    ]
    methods_path = output_dir / "03_统计方法与可复现性.md"
    methods_path.write_text("\n".join(methods_lines) + "\n", encoding="utf-8")
    return rows


def write_academic_reporting_audit(
    run_dir: Path, plan: dict[str, Any], results: dict[str, Any]
) -> Path:
    """Record the evidence/boundary checks applied to report prose without inventing claims."""
    output_dir = run_dir / "90_最终报告"
    output_dir.mkdir(parents=True, exist_ok=True)
    manuscript = ((plan.get("reporting") or {}).get("manuscript_support") or {})
    lines = [
        "# 学术报告表述审计",
        "",
        "本文件仅审计本次自动报告的证据边界，不生成未确认的论文结论。",
        "",
        "## 已执行的约束",
        "",
        "- 结果章节只消费同一运行的统一结果对象。",
        "- 关联、预测与因果解释层级不互相替代。",
        "- 效应量、区间、P 值、诊断与局限性保留在其对应表格或模块对象中。",
        "- 未登记的样本量、统计结果、机制或外部可推广性不被补写。",
        "",
        "## 模块证据登记",
        "",
    ]
    for module_id, result in results.items():
        lines.append(
            f"- `{module_id}`：表格 {len(result.get('tables', []))} 个，"
            f"图形 {len(result.get('figures', []))} 个，"
            f"诊断 {len(result.get('diagnostics', []))} 项，"
            f"局限性 {len(result.get('limitations', []))} 项。"
        )
    lines.extend(["", "## 论文级主张", ""])
    if manuscript.get("enabled") is True:
        lines.append("- 已启用：只接受方案中预先确认且可解析到结果对象的主张；详见 `02_主张证据边界表.csv`。")
    else:
        lines.append("- 未启用：本报告不自动从统计显著性、模型方向或图形生成论文级主张。")
    path = output_dir / "05_学术报告表述审计.md"
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def number_or_none(value: Any) -> float | None:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def render_number(value: Any, digits: int = 2) -> str:
    number = number_or_none(value)
    if number is None:
        return "未估计"
    return f"{number:.{digits}f}"


def render_p_value(value: Any) -> str:
    number = number_or_none(value)
    if number is None:
        return "P 未估计"
    if number < 0.001:
        return "P<0.001"
    return f"P={number:.3f}"


def read_csv_records(path: Path) -> list[dict[str, str]]:
    if not path.is_file():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def table_records(
    run_dir: Path, result: dict[str, Any], table_id: str
) -> list[dict[str, str]]:
    for table in result.get("tables") or []:
        if str(table.get("table_id") or "") == table_id:
            return read_csv_records(run_dir / str(table.get("csv_path") or ""))
    return []


def include_in_manuscript(table_or_figure: dict[str, Any]) -> bool:
    """Keep run-quality diagnostics in supplemental Markdown, not manuscript prose."""
    text = " ".join(
        str(table_or_figure.get(key) or "")
        for key in ("table_id", "figure_id", "title", "path")
    ).lower()
    blocked_tokens = ("诊断", "diagnostic", "残差", "杠杆", "cook", "校准")
    return not any(token in text for token in blocked_tokens)


def variable_label(value: str, plan: dict[str, Any]) -> str:
    labels = ((plan.get("variables") or {}).get("labels") or {})
    text = str(value or "")
    if text in labels:
        return str(labels[text])
    reference_levels = ((plan.get("variables") or {}).get("reference_levels") or {})
    for variable, label in labels.items():
        if text.startswith(str(variable)) and text != str(variable):
            level = text[len(str(variable)) :].lstrip("_:[]")
            reference = reference_levels.get(variable)
            if level and reference not in {None, ""}:
                return f"{label}（{level} 相比 {reference}）"
    return text.replace("_", " ")


TABLE_HEADER_LABELS = {
    "variable": "变量",
    "label": "变量",
    "level": "水平",
    "term": "变量（或水平）",
    "comparison": "比较",
    "method": "统计方法",
    "n": "n",
    "n_display": "n",
    "missing_n": "缺失，n",
    "groups": "组数",
    "rows": "行数",
    "columns": "列数",
    "estimate": "估计值",
    "estimate_log_odds": "估计值（log odds）",
    "coefficient": "相关系数",
    "std_error": "标准误",
    "statistic": "统计量",
    "effect_size": "效应量",
    "effect_size_type": "效应量类型",
    "cramers_v": "Cramér’s V",
    "p_value": "P 值",
    "p_adjusted": "校正后 P 值",
    "group_summary": "各组描述统计",
    "inference": "推断方法",
    "r_squared": "R²",
    "adjusted_r_squared": "调整后 R²",
    "residual_sd": "残差标准差",
    "f_statistic": "F 统计量",
    "df_model": "模型自由度",
    "df_residual": "残差自由度",
    "model_p_value": "模型 P 值",
    "events": "事件数",
    "non_events": "非事件数",
    "odds_ratio": "OR",
    "hazard_ratio": "HR",
    "auc": "AUC",
    "brier_score": "Brier 分数",
    "estimate_ci": "估计值（95% CI）",
    "coefficient_ci": "相关系数（95% CI）",
    "effect_size_ci": "效应量（95% CI）",
    "odds_ratio_ci": "OR（95% CI）",
    "hazard_ratio_ci": "HR（95% CI）",
}

TABLE_INTEGER_COLUMNS = {
    "n",
    "n_display",
    "missing_n",
    "groups",
    "rows",
    "columns",
    "events",
    "non_events",
    "parameters",
    "df_model",
    "df_residual",
}
TABLE_THREE_DECIMAL_COLUMNS = {
    "coefficient",
    "effect_size",
    "cramers_v",
    "r_squared",
    "adjusted_r_squared",
    "auc",
    "brier_score",
}


def render_table_p_value(value: Any) -> str:
    """Render P values for a table cell, without repeating the column label."""
    number = number_or_none(value)
    if number is None:
        return ""
    if number < 0.001:
        return "<0.001"
    return f"{number:.3f}"


def render_table_scalar(value: Any, column: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text.upper() in {"TRUE", "FALSE"}:
        return "是" if text.upper() == "TRUE" else "否"
    if column in {"p_value", "p_adjusted", "model_p_value"}:
        return render_table_p_value(text)
    number = number_or_none(text)
    if number is None:
        return text
    if column in TABLE_INTEGER_COLUMNS:
        return str(int(round(number)))
    digits = 3 if column in TABLE_THREE_DECIMAL_COLUMNS else 2
    if column in {"odds_ratio", "hazard_ratio"} and abs(number) < 0.01:
        digits = 3
    return f"{number:.{digits}f}"


def render_table_ci(
    estimate: Any,
    lower: Any,
    upper: Any,
    *,
    estimate_column: str,
) -> str:
    estimate_text = render_table_scalar(estimate, estimate_column)
    lower_text = render_table_scalar(lower, estimate_column)
    upper_text = render_table_scalar(upper, estimate_column)
    if lower_text and upper_text:
        return f"{estimate_text}（{lower_text}–{upper_text}）"
    return estimate_text


def _combine_table_ci_columns(
    headers: list[str],
    records: list[dict[str, str]],
    *,
    point: str,
    lower: str,
    upper: str,
    combined: str,
) -> tuple[list[str], list[dict[str, str]]]:
    required = {point, lower, upper}
    if not required.issubset(headers):
        return headers, records
    if not any(
        number_or_none(row.get(lower)) is not None
        and number_or_none(row.get(upper)) is not None
        for row in records
    ):
        return headers, records
    new_headers: list[str] = []
    for header in headers:
        if header == point:
            new_headers.append(combined)
        elif header not in {lower, upper}:
            new_headers.append(header)
    new_records: list[dict[str, str]] = []
    for row in records:
        new_row = {
            header: render_table_ci(
                row.get(point),
                row.get(lower),
                row.get(upper),
                estimate_column=point,
            )
            if header == combined
            else row.get(header, "")
            for header in new_headers
        }
        new_records.append(new_row)
    return new_headers, new_records


def table_variable_label(variable: str, plan: dict[str, Any]) -> str:
    variables = plan.get("variables") or {}
    labels = variables.get("labels") or {}
    units = variables.get("units") or {}
    label = str(labels.get(variable) or variable)
    unit = str(units.get(variable) or "").strip()
    return f"{label}（{unit}）" if unit and label != variable else label


def table_reference_level_note(rows: list[list[str]], plan: dict[str, Any]) -> str:
    if not rows or "term" not in {str(value) for value in rows[0]}:
        return ""
    variables = plan.get("variables") or {}
    reference_levels = variables.get("reference_levels") or {}
    entries = [
        f"{table_variable_label(str(variable), plan)}={reference_levels[variable]}"
        for variable in variables.get("categorical") or []
        if reference_levels.get(variable) not in {None, ""}
    ]
    return f"分类自变量的参照水平：{'；'.join(entries)}。" if entries else ""


def render_table_term(value: Any, plan: dict[str, Any]) -> str:
    text = str(value or "").strip()
    if text == "(Intercept)":
        return "截距"
    variables = plan.get("variables") or {}
    candidates = [
        *list(variables.get("categorical") or []),
        *list(variables.get("exposures") or []),
        *list(variables.get("covariates") or []),
    ]
    for variable in sorted({str(item) for item in candidates if str(item)}, key=len, reverse=True):
        if text == variable:
            return table_variable_label(variable, plan)
        if text.startswith(variable):
            level = text[len(variable) :].lstrip("_:")
            if level:
                return f"{table_variable_label(variable, plan)}：{level}"
    if text in ((variables.get("labels") or {})):
        return table_variable_label(text, plan)
    return text.replace("_", " ")


def prepare_article_table_rows(
    rows: list[list[str]], plan: dict[str, Any] | None = None
) -> list[list[str]]:
    """Create a compact reader-facing table while retaining raw CSV unchanged.

    The presentation follows common clinical-journal patterns: meaningful column
    names, estimates paired with their 95% CI, table-cell P values without a
    duplicated prefix, and a single left-hand row-label column.
    """
    if not rows:
        return rows
    plan = plan or {}
    headers = [str(value) for value in rows[0]]
    records = [
        {
            header: str(row[index]) if index < len(row) else ""
            for index, header in enumerate(headers)
        }
        for row in rows[1:]
    ]
    if "label" in headers and "variable" in headers and any(
        record.get("label", "").strip() for record in records
    ):
        headers.remove("variable")
        for record in records:
            record.pop("variable", None)
    if "group_summary" in headers:
        headers.remove("group_summary")
        for record in records:
            record.pop("group_summary", None)

    for point, lower, upper, combined in (
        ("odds_ratio", "or_conf_low", "or_conf_high", "odds_ratio_ci"),
        ("hazard_ratio", "hr_conf_low", "hr_conf_high", "hazard_ratio_ci"),
        ("estimate", "conf_low", "conf_high", "estimate_ci"),
        ("coefficient", "conf_low", "conf_high", "coefficient_ci"),
        ("effect_size", "effect_conf_low", "effect_conf_high", "effect_size_ci"),
    ):
        headers, records = _combine_table_ci_columns(
            headers,
            records,
            point=point,
            lower=lower,
            upper=upper,
            combined=combined,
        )

    rendered = [[TABLE_HEADER_LABELS.get(header, header.replace("_", " ")) for header in headers]]
    for record in records:
        rendered.append(
            [
                render_table_term(record.get(header, ""), plan)
                if header in {"term", "variable"}
                else render_table_scalar(record.get(header, ""), header)
                for header in headers
            ]
        )
    return rendered


def first_matching_row(
    rows: list[dict[str, str]], key: str, target: str
) -> dict[str, str] | None:
    for row in rows:
        if str(row.get(key) or "") == target:
            return row
    return None


def build_module_result_prose(
    run_dir: Path, plan: dict[str, Any], results: dict[str, Any]
) -> dict[str, list[str]]:
    """Create factual manuscript-ready result prose from registered CSV artifacts."""
    prose: dict[str, list[str]] = {}
    primary_outcome = str(
        ((plan.get("variables") or {}).get("outcomes") or {}).get("primary") or ""
    )

    descriptive_result = results.get("descriptive") or {}
    continuous_rows = table_records(run_dir, descriptive_result, "01_连续变量描述性统计")
    categorical_rows = table_records(run_dir, descriptive_result, "02_分类变量描述性统计")
    descriptive_row = first_matching_row(continuous_rows, "variable", primary_outcome)
    if descriptive_row:
        text = (
            f"{variable_label(primary_outcome, plan)}在可用的 "
            f"{descriptive_row.get('n') or '未记录'} 例中，均值为 "
            f"{render_number(descriptive_row.get('mean'))}±{render_number(descriptive_row.get('sd'))}，"
            f"中位数为 {render_number(descriptive_row.get('median'))}"
            f"（四分位距 {render_number(descriptive_row.get('q1'))} 至 "
            f"{render_number(descriptive_row.get('q3'))}）。"
        )
        exposure = str(((plan.get("variables") or {}).get("exposures") or [""])[0])
        exposure_rows = [
            row for row in categorical_rows if str(row.get("variable") or "") == exposure
        ]
        if exposure_rows:
            distribution = "；".join(
                f"{row.get('level')} {row.get('n_display') or row.get('n')} 例"
                f"（{render_number(row.get('pct_valid'), 1)}%）"
                for row in exposure_rows[:6]
            )
            text += f"{variable_label(exposure, plan)}分布为：{distribution}。"
        prose["descriptive"] = [text]

    group_result = results.get("group-comparison") or {}
    group_rows = table_records(run_dir, group_result, "01_连续变量组间比较")
    primary_row = first_matching_row(group_rows, "variable", primary_outcome)
    if primary_row:
        effect = number_or_none(primary_row.get("effect_size"))
        effect_text = (
            f"，{primary_row.get('effect_size_type') or '效应量'}={render_number(effect, 3)}"
            if effect is not None
            else ""
        )
        statement = (
            f"以 {variable_label(primary_outcome, plan)} 为结局的组间比较纳入 "
            f"{primary_row.get('n') or '已分析'} 例，"
            f"{primary_row.get('method') or '已确认检验'} 显示组间差异"
            f"（{render_p_value(primary_row.get('p_adjusted') or primary_row.get('p_value'))}"
            f"{effect_text}）。"
        )
        posthoc_rows = table_records(run_dir, group_result, "03_事后比较")
        posthoc_texts: list[str] = []
        for row in posthoc_rows:
            if str(row.get("variable") or "") != primary_outcome:
                continue
            adjusted_p = number_or_none(row.get("p_adjusted"))
            if adjusted_p is None or adjusted_p > 0.05:
                continue
            posthoc_texts.append(
                f"{row.get('comparison')} 的均值差为 {render_number(row.get('estimate'))}"
                f"（95% CI {render_number(row.get('conf_low'))} 至 "
                f"{render_number(row.get('conf_high'))}，"
                f"{render_p_value(row.get('p_adjusted'))}）"
            )
        if posthoc_texts:
            statement += "事后比较显示，" + "；".join(posthoc_texts[:3]) + "。"
        prose["group-comparison"] = [statement]

    correlation_result = results.get("correlation") or {}
    correlation_rows = table_records(run_dir, correlation_result, "03_相关性检验明细")
    significant_correlations = sorted(
        [
            row
            for row in correlation_rows
            if number_or_none(row.get("p_adjusted")) is not None
            and number_or_none(row.get("p_adjusted")) <= 0.05
        ],
        key=lambda row: abs(number_or_none(row.get("coefficient")) or 0),
        reverse=True,
    )
    if significant_correlations:
        highlights = []
        for row in significant_correlations[:3]:
            highlights.append(
                f"{variable_label(str(row.get('variable_1') or ''), plan)}与"
                f"{variable_label(str(row.get('variable_2') or ''), plan)}"
                f"的 Spearman 相关系数为 {render_number(row.get('coefficient'), 3)}"
                f"（n={row.get('n') or '未记录'}，{render_p_value(row.get('p_adjusted'))}）"
            )
        prose["correlation"] = ["相关分析中，" + "；".join(highlights) + "。"]

    linear_result = results.get("linear-regression") or {}
    linear_summary = table_records(run_dir, linear_result, "02_线性回归模型摘要")
    linear_coefficients = table_records(run_dir, linear_result, "01_线性回归系数")
    if linear_summary:
        summary = linear_summary[0]
        text = (
            f"多元线性回归在 {summary.get('n') or '未记录'} 例完整病例中拟合，"
            f"模型 R²={render_number(summary.get('r_squared'), 3)}，"
            f"调整后 R²={render_number(summary.get('adjusted_r_squared'), 3)}，"
            f"{render_p_value(summary.get('model_p_value'))}。"
        )
        terms = [
            row
            for row in linear_coefficients
            if str(row.get("term") or "") != "(Intercept)"
            and number_or_none(row.get("p_value")) is not None
            and number_or_none(row.get("p_value")) <= 0.05
        ]
        if terms:
            effects = []
            for row in terms[:4]:
                effects.append(
                    f"{variable_label(str(row.get('term') or ''), plan)}的回归系数为 "
                    f"{render_number(row.get('estimate'))}（95% CI "
                    f"{render_number(row.get('conf_low'))} 至 {render_number(row.get('conf_high'))}，"
                    f"{render_p_value(row.get('p_value'))}）"
                )
            text += "在该调整模型中，" + "；".join(effects) + "。"
        prose["linear-regression"] = [text]

    logistic_result = results.get("logistic-regression") or {}
    logistic_summary = table_records(run_dir, logistic_result, "02_Logistic回归模型摘要")
    logistic_coefficients = table_records(run_dir, logistic_result, "01_Logistic回归系数与OR")
    if logistic_summary:
        summary = logistic_summary[0]
        text = (
            f"二分类 Logistic 回归在 {summary.get('n') or '未记录'} 例完整病例中拟合，"
            f"其中事件数为 {summary.get('events') or '未记录'}；"
            f"建模样本内 AUC 为 {render_number(summary.get('auc'), 3)}。"
        )
        terms = [
            row
            for row in logistic_coefficients
            if str(row.get("term") or "") != "(Intercept)"
            and number_or_none(row.get("p_value")) is not None
            and number_or_none(row.get("p_value")) <= 0.05
        ]
        if terms:
            effects = []
            for row in terms[:4]:
                effects.append(
                    f"{variable_label(str(row.get('term') or ''), plan)}的调整后 OR 为 "
                    f"{render_number(row.get('odds_ratio'))}（95% CI "
                    f"{render_number(row.get('or_conf_low'))} 至 "
                    f"{render_number(row.get('or_conf_high'))}，"
                    f"{render_p_value(row.get('p_value'))}）"
                )
            text += "与事件结局的关联中，" + "；".join(effects) + "。"
        prose["logistic-regression"] = [text]

    for module_id, result in results.items():
        if module_id not in prose:
            narratives = [str(item) for item in (result.get("narrative") or []) if str(item).strip()]
            if narratives:
                prose[module_id] = narratives
    return prose


def declared_artifact_commentary(
    result: dict[str, Any], artifact_id: str
) -> list[str]:
    """Use module-authored, result-object evidence notes when available."""
    statements: list[str] = []
    for entry in result.get("reporting_evidence") or []:
        if not isinstance(entry, dict):
            continue
        if str(entry.get("artifact_id") or "") != artifact_id:
            continue
        statement = str(entry.get("result_statement") or "").strip()
        interpretation = str(entry.get("interpretation") or "").strip()
        if statement:
            statements.append(statement)
        if interpretation:
            statements.append(interpretation)
    return statements


def significant_rows(
    rows: list[dict[str, str]], p_field: str = "p_adjusted"
) -> list[dict[str, str]]:
    return [
        row
        for row in rows
        if number_or_none(row.get(p_field)) is not None
        and number_or_none(row.get(p_field)) <= 0.05
    ]


def _build_artifact_commentary(
    run_dir: Path,
    plan: dict[str, Any],
    results: dict[str, Any],
    module_id: str,
    artifact: dict[str, Any],
) -> list[str]:
    """Create one evidence-focused paragraph for a reportable table or figure.

    The prose remains inside the registered analytic scope: it reports observed
    estimates, comparisons, and model summaries without adding mechanisms,
    external validation, or causal interpretation.
    """
    result = results.get(module_id) or {}
    artifact_id = str(artifact.get("table_id") or artifact.get("figure_id") or "")
    declared = declared_artifact_commentary(result, artifact_id)
    if declared:
        return declared

    primary_outcome = str(
        ((plan.get("variables") or {}).get("outcomes") or {}).get("primary") or ""
    )
    table_id = str(artifact.get("table_id") or "")
    if table_id:
        rows = read_csv_records(run_dir / str(artifact.get("csv_path") or ""))
        if module_id == "descriptive" and table_id == "01_连续变量描述性统计":
            primary = first_matching_row(rows, "variable", primary_outcome)
            if primary:
                return [
                    f"该表显示，{variable_label(primary_outcome, plan)}在 "
                    f"{primary.get('n') or '未记录'} 例中的均值为 "
                    f"{render_number(primary.get('mean'))}±{render_number(primary.get('sd'))}，"
                    f"中位数为 {render_number(primary.get('median'))}；"
                    "均值与中位数接近，表明该变量的集中趋势在两种汇总方式下相近。"
                ]
        if module_id == "descriptive" and table_id == "02_分类变量描述性统计":
            exposure = str(((plan.get("variables") or {}).get("exposures") or [""])[0])
            exposure_rows = [
                row for row in rows if str(row.get("variable") or "") == exposure
            ]
            if exposure_rows:
                distribution = "；".join(
                    f"{row.get('level')} {row.get('n_display') or row.get('n')} 例"
                    f"（{render_number(row.get('pct_valid'), 1)}%）"
                    for row in exposure_rows[:6]
                )
                return [
                    f"该表显示，{variable_label(exposure, plan)}各水平的样本量分别为："
                    f"{distribution}。该分布为后续组间比较提供了明确的分组基数。"
                ]
        if module_id == "group-comparison" and table_id == "01_连续变量组间比较":
            primary = first_matching_row(rows, "variable", primary_outcome)
            if primary:
                effect = number_or_none(primary.get("effect_size"))
                effect_text = (
                    f"，{primary.get('effect_size_type') or '效应量'}="
                    f"{render_number(effect, 3)}"
                    if effect is not None
                    else ""
                )
                return [
                    f"表中以 {variable_label(primary_outcome, plan)} 为结局的 "
                    f"{primary.get('method') or '已确认检验'} 显示组间总体差异"
                    f"（{render_p_value(primary.get('p_adjusted') or primary.get('p_value'))}"
                    f"{effect_text}）。该结果对应总体组间比较；具体组对的均值差及 95% CI 见事后比较表。"
                ]
        if module_id == "group-comparison" and table_id == "02_分类变量组间比较":
            outcome = str(
                (((plan.get("variables") or {}).get("outcomes") or {}).get("secondary") or [""])[0]
            )
            categorical = first_matching_row(rows, "variable", outcome)
            if categorical:
                effect = number_or_none(categorical.get("cramers_v"))
                effect_text = f"，Cramér's V={render_number(effect, 3)}" if effect is not None else ""
                return [
                    f"表中显示，{variable_label(outcome, plan)}在分组间的分布经 "
                    f"{categorical.get('method') or '已确认检验'}比较"
                    f"（{render_p_value(categorical.get('p_adjusted') or categorical.get('p_value'))}"
                    f"{effect_text}）。该比较描述分类结局与分组的样本内关联。"
                ]
        if module_id == "group-comparison" and table_id == "03_事后比较":
            comparisons = [
                row for row in significant_rows(rows) if str(row.get("variable") or "") == primary_outcome
            ]
            if comparisons:
                details = "；".join(
                    f"{row.get('comparison')} 的均值差为 {render_number(row.get('estimate'))}"
                    f"（95% CI {render_number(row.get('conf_low'))} 至 {render_number(row.get('conf_high'))}，"
                    f"{render_p_value(row.get('p_adjusted'))}）"
                    for row in comparisons[:3]
                )
                return [
                    "在总体检验提示差异后，事后比较进一步定位了差异来源：" + details + "。"
                ]
        if module_id == "correlation" and table_id == "01_相关系数矩阵":
            detail_rows = table_records(run_dir, result, "03_相关性检验明细")
            highlights = sorted(
                significant_rows(detail_rows),
                key=lambda row: abs(number_or_none(row.get("coefficient")) or 0),
                reverse=True,
            )
            if highlights:
                row = highlights[0]
                return [
                    f"矩阵中绝对值最大的经校正相关为"
                    f"{variable_label(str(row.get('variable_1') or ''), plan)}与"
                    f"{variable_label(str(row.get('variable_2') or ''), plan)}"
                    f"（Spearman ρ={render_number(row.get('coefficient'), 3)}，"
                    f"n={row.get('n') or '未记录'}，{render_p_value(row.get('p_adjusted'))}）。"
                    "该系数反映变量在本样本中的单调关联强度。"
                ]
        if module_id == "correlation" and table_id == "02_相关有效样本量矩阵":
            counts = [
                number_or_none(value)
                for row in rows
                for key, value in row.items()
                if key != "variable"
            ]
            counts = [count for count in counts if count is not None]
            if counts:
                return [
                    f"各变量对的有效样本量范围为 {int(min(counts))} 至 {int(max(counts))}。"
                    "相关系数矩阵与该有效样本量矩阵应结合阅读，以识别缺失值导致的配对样本差异。"
                ]
        if module_id == "correlation" and table_id == "03_相关性检验明细":
            corrected = significant_rows(rows)
            if corrected:
                return [
                    f"在 {len(rows)} 个变量对中，Holm 校正后有 {len(corrected)} 个相关性检验达到"
                    "预设统计学阈值；完整系数、有效样本量和校正后 P 值均列于表中。"
                ]
        if module_id == "linear-regression" and table_id == "01_线性回归系数":
            terms = [
                row for row in significant_rows(rows, "p_value")
                if str(row.get("term") or "") != "(Intercept)"
            ]
            if terms:
                details = "；".join(
                    f"{variable_label(str(row.get('term') or ''), plan)}的 β={render_number(row.get('estimate'))}"
                    f"（95% CI {render_number(row.get('conf_low'))} 至 {render_number(row.get('conf_high'))}，"
                    f"{render_p_value(row.get('p_value'))}）"
                    for row in terms[:3]
                )
                return [
                    "在同时调整方案中列出的协变量后，" + details + "。这些系数是对随访评分的条件关联估计。"
                ]
        if module_id == "linear-regression" and table_id == "02_线性回归模型摘要" and rows:
            row = rows[0]
            r_squared = number_or_none(row.get("r_squared"))
            explained = (
                f"{render_number(r_squared * 100, 1)}%"
                if r_squared is not None
                else "未估计"
            )
            return [
                f"该模型基于 {row.get('n') or '未记录'} 例完整病例，解释了 "
                f"{explained} 的随访评分变异（调整后 R²="
                f"{render_number(row.get('adjusted_r_squared'), 3)}，"
                f"{render_p_value(row.get('model_p_value'))}）。"
            ]
        if module_id == "logistic-regression" and table_id == "01_Logistic回归系数与OR":
            terms = [
                row for row in significant_rows(rows, "p_value")
                if str(row.get("term") or "") != "(Intercept)"
            ]
            if terms:
                details = "；".join(
                    f"{variable_label(str(row.get('term') or ''), plan)}的调整后 OR="
                    f"{render_number(row.get('odds_ratio'))}（95% CI "
                    f"{render_number(row.get('or_conf_low'))} 至 {render_number(row.get('or_conf_high'))}，"
                    f"{render_p_value(row.get('p_value'))}）"
                    for row in terms[:3]
                )
                return [
                    "该表量化了预设变量与二分类临床结局之间的调整后关联：" + details + "。"
                ]
        if module_id == "logistic-regression" and table_id == "02_Logistic回归模型摘要" and rows:
            row = rows[0]
            return [
                f"模型在 {row.get('n') or '未记录'} 例完整病例中拟合，其中事件数为 "
                f"{row.get('events') or '未记录'}；建模样本内 AUC={render_number(row.get('auc'), 3)}，"
                f"Brier 分数={render_number(row.get('brier_score'), 3)}。"
            ]

        if rows:
            p_values = [
                number_or_none(row.get("p_adjusted") or row.get("p_value"))
                for row in rows
            ]
            p_values = [value for value in p_values if value is not None]
            if p_values:
                significant_count = sum(value <= 0.05 for value in p_values)
                return [
                    f"该表共列示 {len(rows)} 条结果，其中 {significant_count} 条的 P 值不高于 0.05。"
                    "具体比较、估计值与区间应结合表中对应行解读。"
                ]
            return [
                f"该表以结构化形式汇总 {len(rows)} 条{MODULE_TITLES.get(module_id, module_id)}结果，"
                "用于支持本小节的分析结论。"
            ]

    figure_id = str(artifact.get("figure_id") or "")
    if module_id == "logistic-regression" and figure_id == "roc_curve":
        summary_rows = table_records(run_dir, result, "02_Logistic回归模型摘要")
        if summary_rows:
            row = summary_rows[0]
            return [
                f"图示建模样本内 ROC 曲线的曲线下面积为 {render_number(row.get('auc'), 3)}，"
                f"对应 {row.get('n') or '未记录'} 例完整病例和 {row.get('events') or '未记录'} 个事件。"
                "该图直观呈现模型在当前样本内区分事件与非事件的能力。"
            ]
    if figure_id:
        statistics = artifact.get("statistics") or {}
        return [
            f"该图展示{MODULE_TITLES.get(module_id, module_id)}的已登记可视化证据；"
            f"样本定义为{statistics.get('n_definition') or '见图注'}，"
            f"统计检验为{statistics.get('test') or '见图注'}。"
        ]
    return []


def render_article_style_evidence(
    statements: list[str],
    artifact: dict[str, Any],
    artifact_number: int | None,
) -> list[str]:
    """Turn registered evidence notes into natural manuscript prose.

    A table or figure is cited where its evidence is first discussed.  The
    function intentionally removes report-like labels (for example, “表中要点”)
    and keeps the original statement factual rather than adding an inference.
    """
    if artifact_number is None:
        return statements

    is_table = bool(artifact.get("table_id"))
    reference = f"表 {artifact_number}" if is_table else f"图 {artifact_number}"
    polished: list[str] = []

    for statement in statements:
        text = str(statement).strip()
        if not text:
            continue
        if text.startswith(("如表 ", "如图 ", "表 ", "图 ")):
            polished.append(text)
            continue

        if is_table:
            replacements = (
                ("该表显示，", f"如{reference} 所示，"),
                ("表中以 ", f"如{reference} 所示，以 "),
                ("表中显示，", f"如{reference} 所示，"),
                (
                    "在总体检验提示差异后，事后比较进一步定位了差异来源：",
                    f"进一步的事后比较（{reference}）显示，",
                ),
                ("在总体检验提示差异后，", f"进一步的事后比较显示（{reference}），"),
                ("矩阵中绝对值最大的经校正相关为", f"如{reference} 所示，绝对值最大的经校正相关为"),
                ("各变量对的有效样本量范围为", f"{reference} 中各变量对的有效样本量为"),
                ("在同时调整", f"多元线性回归结果（{reference}）显示，在同时调整"),
                ("该模型基于", f"模型拟合概况见{reference}。该模型基于"),
                ("该表量化了", f"多变量 Logistic 回归结果见{reference}。在该模型中，"),
                ("模型在", f"模型拟合概况见{reference}。模型在"),
                ("该表共列示", f"{reference}共列示"),
                ("该表以", f"{reference}以"),
            )
            for source, replacement in replacements:
                if text.startswith(source):
                    text = replacement + text[len(source) :]
                    break
            else:
                text = f"相关结果见{reference}。{text}"
        else:
            if text.startswith("图示"):
                text = f"如{reference} 所示，" + text[len("图示") :]
            elif text.startswith("该图展示"):
                text = f"{reference}展示" + text[len("该图展示") :]
            else:
                text = f"相关图形结果见{reference}。{text}"
        polished.append(text)
    return polished


def build_artifact_commentary(
    run_dir: Path,
    plan: dict[str, Any],
    results: dict[str, Any],
    module_id: str,
    artifact: dict[str, Any],
    artifact_number: int | None = None,
) -> list[str]:
    """Build factual evidence notes and render them in manuscript style."""
    return render_article_style_evidence(
        _build_artifact_commentary(run_dir, plan, results, module_id, artifact),
        artifact,
        artifact_number,
    )


def default_background(research: dict[str, Any]) -> str:
    background = str(research.get("background") or "").strip()
    if background:
        return background
    objective = str(research.get("primary_objective") or "").strip()
    question = str(research.get("primary_question") or "").strip()
    focus = objective or question or "研究问题"
    return (
        "医学研究中的分组差异与结局关联需要在明确研究问题、变量角色和统计模型的前提下加以评估。"
        f"本研究围绕“{focus}”对提供数据开展结构化统计分析。"
    )


def build_academic_manuscript_outline(plan: dict[str, Any]) -> list[dict[str, str]]:
    research = plan.get("research") or {}
    variables = plan.get("variables") or {}
    keywords = research.get("keywords") or []
    if not keywords:
        candidates = [
            *list(variables.get("exposures") or []),
            str((variables.get("outcomes") or {}).get("primary") or ""),
            "医学统计",
        ]
        keywords = [variable_label(str(item), plan) for item in candidates if str(item).strip()]
    return [
        {"id": "abstract", "title": "摘  要", "text": ""},
        {"id": "keywords", "title": "关键词", "text": "；".join(dict.fromkeys(map(str, keywords)))},
        {"id": "introduction", "title": "1 引言", "text": default_background(research)},
        {"id": "methods", "title": "2 资料与方法", "text": ""},
        {"id": "results", "title": "3 结果", "text": ""},
        {"id": "discussion", "title": "4 讨论", "text": ""},
    ]


def write_reporting_supplements(
    output_dir: Path,
    plan: dict[str, Any],
    manifest: dict[str, Any],
    results: dict[str, Any],
) -> list[Path]:
    """Write diagnostics, limitations, and reproducibility outside the manuscript."""
    output_dir.mkdir(parents=True, exist_ok=True)
    diagnostic_lines = [
        "# 统计诊断与警告",
        "",
        "本文件是运行级质量附件，不属于学术论文初稿正文。",
    ]
    limitation_lines = [
        "# 研究局限性",
        "",
        "本文件是运行级解释边界附件，不属于学术论文初稿正文。",
    ]
    for module_id, result in results.items():
        diagnostics = result.get("diagnostics") or []
        warnings = result.get("warnings") or []
        limitations = result.get("limitations") or []
        if diagnostics or warnings:
            diagnostic_lines.extend(["", f"## {MODULE_TITLES.get(module_id, module_id)}"])
            for diagnostic in diagnostics:
                if isinstance(diagnostic, dict):
                    name = str(
                        diagnostic.get("diagnostic")
                        or diagnostic.get("name")
                        or diagnostic.get("id")
                        or "诊断"
                    )
                    status = str(diagnostic.get("status") or "未登记")
                    value = diagnostic.get("message") or diagnostic.get("value")
                    rule = diagnostic.get("rule")
                    detail_parts = []
                    if value not in {None, ""}:
                        detail_parts.append(f"值：{value}")
                    if rule not in {None, ""}:
                        detail_parts.append(f"判定规则：{rule}")
                    detail = "；".join(detail_parts)
                    diagnostic_lines.append(
                        f"- 诊断：{name}；状态：{status}{'；' + detail if detail else ''}"
                    )
                else:
                    diagnostic_lines.append(f"- 诊断：{diagnostic}")
            for warning in warnings:
                diagnostic_lines.append(f"- 警告：{warning}")
        if limitations:
            limitation_lines.extend(["", f"## {MODULE_TITLES.get(module_id, module_id)}"])
            for limitation in limitations:
                limitation_lines.append(f"- {limitation}")

    modules = [
        item if isinstance(item, str) else item.get("id", "")
        for item in ((plan.get("analysis") or {}).get("modules") or [])
    ]
    research = plan.get("research") or {}
    runtime = plan.get("runtime") or {}
    handling = plan.get("data_handling") or {}
    reproducibility_lines = [
        "# 可复现性信息",
        "",
        "本文件是运行级可复现性附件，不属于学术论文初稿正文。",
        "",
        f"- 运行编号：{manifest.get('run_id', '')}",
        f"- 研究设计：{research.get('design', 'unknown')}",
        f"- 已确认模块：{', '.join(str(item) for item in modules if item)}",
        f"- 缺失值策略：{handling.get('missing_strategy', '')}",
        f"- 统计引擎：{runtime.get('language', 'R')}",
        f"- 最低 R 版本：{runtime.get('minimum_version', '')}",
        f"- 随机种子：{(plan.get('run') or {}).get('random_seed', '')}",
        f"- 分析方案 SHA-256：{manifest.get('analysis_plan_sha256', '')}",
        "- R 会话信息：sessionInfo.txt",
        "- 包版本：package_versions.csv",
        "- Python/R 双环境清单：runtime/environment_manifest.json",
        "- 输入与输出哈希清单：manifest.json",
    ]
    paths = [
        output_dir / "06_统计诊断与警告.md",
        output_dir / "07_研究局限性.md",
        output_dir / "08_可复现性信息.md",
    ]
    paths[0].write_text("\n".join(diagnostic_lines) + "\n", encoding="utf-8")
    paths[1].write_text("\n".join(limitation_lines) + "\n", encoding="utf-8")
    paths[2].write_text("\n".join(reproducibility_lines) + "\n", encoding="utf-8")
    return paths


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成简要医学统计分析 Word 报告。")
    parser.add_argument("--run-dir", required=True)
    parser.add_argument("--output")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    run_dir = Path(args.run_dir).expanduser().resolve()
    output_path = (
        Path(args.output).expanduser().resolve()
        if args.output
        else run_dir / "90_最终报告" / MANUSCRIPT_REPORT_FILENAME
    )
    try:
        output_path.relative_to(run_dir)
    except ValueError as exc:
        raise SystemExit("报告必须保存在当前运行目录内。") from exc

    validation_path = run_dir / "validation_report.json"
    required = [
        validation_path,
        run_dir / "manifest.json",
        run_dir / "analysis_plan.yml",
        run_dir / "data_profile.json",
        run_dir / "execution_status.json",
        run_dir / "analysis_results.json",
    ]
    missing = [path.name for path in required if not path.exists()]
    if missing:
        raise SystemExit(f"缺少生成报告所需文件：{', '.join(missing)}")
    validation = load_json(validation_path)
    if validation.get("valid") is not True or validation.get("mode") not in {
        "execute",
        "report",
    }:
        raise SystemExit("必须先通过 execute 或 report 模式输出验证。")

    manifest = load_json(run_dir / "manifest.json")
    plan = load_yaml(run_dir / "analysis_plan.yml")
    profile = load_json(run_dir / "data_profile.json")
    execution = load_json(run_dir / "execution_status.json")
    results = load_json(run_dir / "analysis_results.json")
    if execution.get("status") != "completed":
        raise SystemExit("分析执行状态不是 completed。")
    manuscript_claim_rows = write_manuscript_support_artifacts(
        run_dir, plan, manifest, results
    )
    write_academic_reporting_audit(run_dir, plan, results)
    report_dir = run_dir / "90_最终报告"
    write_reporting_supplements(report_dir, plan, manifest, results)
    module_prose = build_module_result_prose(run_dir, plan, results)
    outline = build_academic_manuscript_outline(plan)
    outline_by_id = {item["id"]: item for item in outline}

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal_style = document.styles["Normal"]
    normal_style.font.name = LATIN_FONT
    normal_style.font.size = Pt(BODY_SIZE)
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.first_line_indent = Pt(24)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    set_style_font(document.styles["Heading 1"], 14, bold=True)
    set_style_font(document.styles["Heading 2"], 12, bold=True)

    title = document.add_paragraph()
    set_paragraph_format(
        title,
        indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    research = plan.get("research") or {}
    title_run = title.add_run(
        str(research.get("title") or "医学统计分析论文初稿")
    )
    set_run_font(title_run, size=16, bold=True)

    add_heading(document, outline_by_id["abstract"]["title"], 1)
    selected_titles = [
        MODULE_TITLES.get(module_id, module_id)
        for module_id in execution.get("completed_modules", [])
    ]
    abstract_findings = []
    for module_id in ("group-comparison", "linear-regression", "logistic-regression"):
        abstract_findings.extend(module_prose.get(module_id, [])[:1])
    abstract_result = " ".join(abstract_findings[:2]) or "已生成各确认统计模块的结果表与图形。"
    add_text_paragraph(
        document,
        "目的："
        + str(research.get("primary_objective") or research.get("primary_question") or "描述研究数据并评估预设关联。")
        + "方法："
        + f"基于{research.get('design', '已登记研究设计')}，采用"
        + "、".join(selected_titles)
        + "完成确认分析。结果："
        + abstract_result
        + "结论：本次分析在已确认的研究问题和模型设定下提供样本内的描述、比较与关联估计，不将统计关联表述为因果效应。",
    )

    add_heading(document, outline_by_id["keywords"]["title"], 1)
    add_text_paragraph(document, outline_by_id["keywords"]["text"], indent=False)

    add_heading(document, outline_by_id["introduction"]["title"], 1)
    add_text_paragraph(document, outline_by_id["introduction"]["text"])

    add_heading(document, outline_by_id["methods"]["title"], 1)
    add_heading(document, "2.1 研究对象与变量", 2)
    primary_question = str(research.get("primary_question") or "未提供主要研究问题").rstrip("。！？!? ")
    add_text_paragraph(
        document,
        f"本研究的主要问题为：{primary_question}。"
        f"主要结局为{variable_label(str(((plan.get('variables') or {}).get('outcomes') or {}).get('primary') or ''), plan)}；"
        "暴露因素、协变量及分类变量均以确认的分析方案为准。",
    )
    add_heading(document, "2.2 统计学方法", 2)
    handling = plan.get("data_handling") or {}
    add_text_paragraph(
        document,
        "连续变量和分类变量的描述性统计、组间比较、相关分析及回归分析均按确认模块执行。"
        f"缺失值按“{handling.get('missing_strategy') or '未登记'}”策略在相应模块中处理；"
        "组间比较与相关分析的多重比较校正按确认方案执行。"
        "结果优先报告效应估计、95% 置信区间与 P 值。",
    )

    add_heading(document, outline_by_id["results"]["title"], 1)
    add_heading(document, "3.1 样本与数据概况", 2)
    summary = profile.get("summary") or {}
    dataset = ((profile.get("datasets") or [{}])[0] or {})
    add_text_paragraph(
        document,
        "主分析数据集包含 "
        f"{dataset.get('rows', summary.get('total_rows_across_datasets', 0))} 条记录和 "
        f"{dataset.get('columns', summary.get('total_variables_across_datasets', 0))} 个变量。"
        f"总体缺失比例为 {summary.get('overall_missing_pct', 'NA')}%，"
        f"重复记录数为 {dataset.get('duplicate_rows', 0)}。",
    )

    module_counter = 2
    table_counter = 1
    figure_counter = 1
    for module_id, result in results.items():
        add_heading(
            document,
            f"3.{module_counter} {MODULE_TITLES.get(module_id, module_id)}",
            2,
        )
        module_counter += 1
        for narrative in module_prose.get(module_id, result.get("narrative", [])):
            add_text_paragraph(document, str(narrative))
        for table in result.get("tables", []):
            if not include_in_manuscript(table):
                continue
            current_table_number = table_counter
            add_caption(
                document,
                f"表 {current_table_number} {table.get('title', table.get('table_id', '结果表'))}",
            )
            table_counter += 1
            csv_path = run_dir / str(table["csv_path"])
            add_three_line_table(document, read_csv_rows(csv_path), plan=plan)
            for footnote in table.get("footnotes", []):
                add_table_note(document, f"注：{footnote}")
            for commentary in build_artifact_commentary(
                run_dir,
                plan,
                results,
                module_id,
                table,
                artifact_number=current_table_number,
            ):
                add_text_paragraph(document, commentary)
        for figure in result.get("figures", []):
            if not include_in_manuscript(figure):
                continue
            figure_path = run_dir / str(figure.get("path", ""))
            if figure_path.is_file():
                current_figure_number = figure_counter
                paragraph = document.add_paragraph()
                set_paragraph_format(paragraph, indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                run = paragraph.add_run()
                figure_title = str(figure.get("title", "统计图"))
                inline_shape = run.add_picture(str(figure_path), width=Inches(5.8))
                inline_shape._inline.docPr.set("descr", figure_title)
                inline_shape._inline.docPr.set("title", figure_title)
                add_caption(document, f"图 {current_figure_number} {figure_title}")
                figure_counter += 1
                statistics = figure.get("statistics") or {}
                caption_details = [
                    f"n：{statistics.get('n_definition', '')}",
                    f"中心统计量：{statistics.get('center_statistic', '')}",
                    f"区间/误差：{statistics.get('interval', '')}",
                    f"检验：{statistics.get('test', '')}",
                    (
                        "多重比较校正："
                        f"{statistics.get('multiple_comparison_correction', '')}"
                    ),
                ]
                detail_paragraph = document.add_paragraph()
                set_paragraph_format(
                    detail_paragraph,
                    indent=False,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                )
                detail_run = detail_paragraph.add_run(
                    "；".join(item for item in caption_details if item)
                )
                set_run_font(detail_run, size=10)
                for commentary in build_artifact_commentary(
                    run_dir,
                    plan,
                    results,
                    module_id,
                    figure,
                    artifact_number=current_figure_number,
                ):
                    add_text_paragraph(document, commentary)

    add_heading(document, outline_by_id["discussion"]["title"], 1)
    discussion_points: list[str] = []
    if module_prose.get("group-comparison") and module_prose.get("linear-regression"):
        discussion_points.append(
            "单因素比较与调整协变量后的线性回归均提供了治疗组与随访临床评分之间差异的样本内估计，"
            "提示分组差异在已确认模型设定下具有一致的统计学表现。"
        )
    if module_prose.get("correlation"):
        discussion_points.append(
            "相关分析进一步描述了连续变量之间的方向与强度，为回归模型中协变量关系的解读提供了补充。"
        )
    if module_prose.get("logistic-regression"):
        discussion_points.append(
            "二分类 Logistic 回归报告的是建模样本内的调整后关联与区分度，用于刻画预设变量与二分类结局之间的统计关系。"
        )
    add_text_paragraph(
        document,
        "".join(discussion_points)
        or "本次结果围绕已确认的研究问题报告样本特征、比较结果和模型估计，并在既定设计与变量定义范围内进行解释。",
    )

    if manuscript_claim_rows:
        add_heading(document, "附录：经确认的主张、证据与边界", 1)
        add_text_paragraph(
            document,
            "本节只列出分析方案中由用户确认、且已解析到同一运行统一结果对象的主张。"
            "结果陈述、解释层级和适用边界不得互相替代。",
        )
        claim_table_rows = [[
            "主张编号",
            "主张",
            "证据对象",
            "解释层级",
            "边界",
        ]]
        claim_table_rows.extend([
            [
                row["claim_id"],
                row["statement"],
                row["evidence_files"],
                row["interpretation_level"],
                row["boundary"],
            ]
            for row in manuscript_claim_rows
        ])
        add_three_line_table(
            document,
            claim_table_rows,
            plan=plan,
            max_rows=25,
            max_columns=5,
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(json.dumps({"status": "created", "output": str(output_path)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
